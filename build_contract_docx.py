from pathlib import Path
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "Rotsi_Labs_Website_Services_Agreement_Template.md"
OUTPUT = ROOT / "Rotsi_Labs_Neighbourhood_Cocktails_Website_Services_Agreement.docx"
DOCS_SKILL = Path(r"C:\Users\hp\.codex\plugins\cache\openai-primary-runtime\documents\26.430.10722\skills\documents")
sys.path.insert(0, str(DOCS_SKILL / "scripts"))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


ACCENT = "1F4E5F"
MUTED = "5F6B73"
LIGHT = "EAF0F2"
BORDER = "C9D2D6"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=BORDER, size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_header_repeat(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_border(paragraph, color=ACCENT):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), ACCENT)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_markdown_runs(paragraph, text):
    pattern = re.compile(r"\*\*(.+?)\*\*|\[(.+?)\]\((.+?)\)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        if match.group(1) is not None:
            run = paragraph.add_run(match.group(1))
            run.bold = True
        else:
            add_hyperlink(paragraph, match.group(2), match.group(3))
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def clean_text(text):
    return text.replace("**", "").strip()


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(32, 36, 38)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, before, after in (
        ("Title", 22, 0, 8),
        ("Subtitle", 11, 0, 10),
        ("Heading 1", 15, 14, 6),
        ("Heading 2", 12.5, 10, 4),
        ("Heading 3", 11, 8, 3),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(ACCENT if name != "Subtitle" else MUTED)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = name.startswith("Heading")

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.08

    if "Contract Metadata" not in styles:
        style = styles.add_style("Contract Metadata", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(9.5)
        style.font.color.rgb = RGBColor.from_string(MUTED)


def setup_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "Rotsi Labs / Neighbourhood Cocktails - Website Services Agreement"
    hp.style = "Contract Metadata"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_border(hp)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.style = "Contract Metadata"
    fp.add_run("Page ")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = fp.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_metadata_table(doc):
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    rows = [
        ("Service Provider", "Rotsi Api Solutions, trading as Rotsi Labs"),
        ("Client", "[Client legal name], trading as Neighbourhood Cocktails"),
        ("Project", "Landing-page website, embedded Substack integration, booking data collection, deployment, and downtime maintenance"),
        ("Start / Duration", "18 May 2026 / maximum 3 weeks, subject to permitted extensions"),
    ]
    for r_idx, (label, value) in enumerate(rows):
        table.cell(r_idx, 0).text = label
        table.cell(r_idx, 1).text = value
    apply_table_geometry(table, [2200, 7160], table_width_dxa=9360, indent_dxa=0)
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            set_cell_borders(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.style = "Normal"
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
            if c_idx == 0:
                set_cell_shading(cell, LIGHT)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(ACCENT)
            else:
                set_cell_shading(cell, "FFFFFF")
    set_table_header_repeat(table.rows[0])
    doc.add_paragraph()


def extract_contract(markdown):
    start = markdown.index("### Website Design, Development and Maintenance Services Agreement")
    end = markdown.index("## 4. Lawyer Verification Note")
    return markdown[start:end].strip()


def split_table(lines, start):
    table_lines = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        table_lines.append(lines[i].strip())
        i += 1
    rows = []
    for line in table_lines:
        parts = [cell.strip() for cell in line.strip("|").split("|")]
        if all(set(cell) <= {"-", ":", " "} for cell in parts):
            continue
        rows.append(parts)
    return rows, i


def table_widths(headers):
    count = len(headers)
    joined = " ".join(headers).lower()
    if count == 4:
        return column_widths_from_weights([1.2, 3.4, 2.6, 2.2], 9360)
    if count == 3 and "amount" in joined:
        return column_widths_from_weights([4.1, 2.1, 3.2], 9360)
    if count == 3:
        return column_widths_from_weights([3.0, 3.0, 3.4], 9360)
    return column_widths_from_weights([1] * count, 9360)


def add_markdown_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.style = "Normal"
            add_markdown_runs(paragraph, value)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if r_idx == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(ACCENT)
            set_cell_borders(cell)
            set_cell_margins(cell, top=120, bottom=120, start=120, end=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, LIGHT)
            else:
                set_cell_shading(cell, "FFFFFF")
    set_table_header_repeat(table.rows[0])
    widths = table_widths(rows[0])
    apply_table_geometry(table, widths, table_width_dxa=sum(widths), indent_dxa=0)
    doc.add_paragraph()


def add_signature_line(doc, line):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    add_markdown_runs(p, line)


def build_doc():
    markdown = SOURCE.read_text(encoding="utf-8")
    contract = extract_contract(markdown)
    lines = contract.splitlines()

    doc = Document()
    configure_styles(doc)
    setup_section(doc.sections[0])

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Website Design, Development and Maintenance Services Agreement")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Clean signing draft | Rotsi Labs and Neighbourhood Cocktails")

    add_metadata_table(doc)

    p = doc.add_paragraph(style="Contract Metadata")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Prepared from the working contract draft. Bracketed fields remain editable before signing.")

    doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(doc.sections[-1])

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue

        if line.startswith("|"):
            rows, i = split_table(lines, i)
            add_markdown_table(doc, rows)
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            text = clean_text(heading.group(2))
            if text == "Website Design, Development and Maintenance Services Agreement":
                i += 1
                continue
            style = "Heading 1" if level <= 2 else "Heading 2"
            if level >= 4:
                style = "Heading 3"
            p = doc.add_paragraph(style=style)
            p.add_run(text)
            i += 1
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_markdown_runs(p, line[2:].strip())
            i += 1
            continue

        if re.match(r"^[A-Za-z ]+:\s*_+", line) or re.match(r"^(Name|Title|Signature|Date|Witness)", line):
            add_signature_line(doc, line)
            i += 1
            continue

        p = doc.add_paragraph()
        add_markdown_runs(p, line)
        i += 1

    doc.core_properties.title = "Website Design, Development and Maintenance Services Agreement"
    doc.core_properties.subject = "Rotsi Labs and Neighbourhood Cocktails"
    doc.core_properties.author = "Rotsi Labs"
    doc.core_properties.comments = "Draft agreement generated from reviewed Markdown template."
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
